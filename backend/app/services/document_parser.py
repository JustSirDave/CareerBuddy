"""
CareerBuddy - Document Parser Service
Extracts content from uploaded resume files (DOCX, PDF)
Author: Sir Dave
"""
from pathlib import Path
from typing import Dict, Optional
from loguru import logger
import re


def extract_from_docx(file_path: Path) -> str:
    """
    Extract text content from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        
        # Extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in paragraphs:  # Avoid duplicates
                        paragraphs.append(text)
        
        content = '\n'.join(paragraphs)
        logger.info(f"[extract_from_docx] Extracted {len(content)} characters from {file_path.name}")
        
        return content
        
    except Exception as e:
        logger.error(f"[extract_from_docx] Failed to extract from {file_path}: {e}")
        raise Exception(f"Unable to read DOCX file: {str(e)}")


def extract_from_pdf(file_path: Path) -> str:
    """
    Extract text content from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import pdfplumber
        
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    text_content.append(text.strip())
                    logger.debug(f"[extract_from_pdf] Extracted page {page_num}: {len(text)} chars")
        
        content = '\n\n'.join(text_content)
        logger.info(f"[extract_from_pdf] Extracted {len(content)} characters from {file_path.name} ({len(pdf.pages)} pages)")
        
        return content
        
    except Exception as e:
        logger.error(f"[extract_from_pdf] Failed to extract from {file_path}: {e}")
        # Try fallback to pypdf
        try:
            return extract_from_pdf_fallback(file_path)
        except:
            raise Exception(f"Unable to read PDF file: {str(e)}")


def extract_from_pdf_fallback(file_path: Path) -> str:
    """
    Fallback PDF extraction using pypdf.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_content = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                text_content.append(text.strip())
                logger.debug(f"[extract_from_pdf_fallback] Extracted page {page_num}: {len(text)} chars")
        
        content = '\n\n'.join(text_content)
        logger.info(f"[extract_from_pdf_fallback] Extracted {len(content)} characters from {file_path.name}")
        
        return content
        
    except Exception as e:
        logger.error(f"[extract_from_pdf_fallback] Failed: {e}")
        raise Exception(f"Unable to read PDF file: {str(e)}")


def parse_document(file_path: Path, file_type: str) -> Dict[str, any]:
    """
    Parse uploaded document and extract content.
    
    Args:
        file_path: Path to uploaded file
        file_type: File type (docx, pdf)
        
    Returns:
        Dictionary with:
        - content: Extracted text
        - sections: Detected sections (if any)
        - word_count: Number of words
        - char_count: Number of characters
    """
    try:
        # Extract content based on file type
        if file_type == 'docx':
            content = extract_from_docx(file_path)
        elif file_type == 'pdf':
            content = extract_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Basic validation
        if not content or len(content.strip()) < 50:
            raise ValueError("Document appears to be empty or too short. Please upload a valid resume.")
        
        # Detect sections (basic pattern matching)
        sections = detect_sections(content)
        
        # Calculate stats
        word_count = len(content.split())
        char_count = len(content)
        
        logger.info(f"[parse_document] Parsed {file_path.name}: {word_count} words, {len(sections)} sections")
        
        return {
            'content': content,
            'sections': sections,
            'word_count': word_count,
            'char_count': char_count,
            'file_type': file_type,
            'file_name': file_path.name
        }
        
    except Exception as e:
        logger.error(f"[parse_document] Error parsing {file_path}: {e}")
        raise


def detect_sections(content: str) -> Dict[str, str]:
    """
    Detect common resume sections in the content.
    
    Args:
        content: Resume text content
        
    Returns:
        Dictionary of detected sections
    """
    sections = {}
    
    # Common section headers (case-insensitive)
    section_patterns = {
        'contact': r'(contact|personal details|personal information)',
        'summary': r'(summary|profile|objective|about me)',
        'experience': r'(experience|work history|employment|professional experience)',
        'education': r'(education|academic|qualifications)',
        'skills': r'(skills|technical skills|core competencies|expertise)',
        'certifications': r'(certifications?|licenses?|credentials)',
        'projects': r'(projects?|portfolio)',
        'achievements': r'(achievements?|awards?|honors?)',
    }
    
    for section_name, pattern in section_patterns.items():
        match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
        if match:
            sections[section_name] = match.group(0)
            logger.debug(f"[detect_sections] Found section: {section_name}")
    
    return sections


def validate_file_format(file_name: str, mime_type: str, user_tier: str = "free") -> tuple[bool, str, str]:
    """
    Validate uploaded file format based on user tier.
    
    Args:
        file_name: Name of uploaded file
        mime_type: MIME type from Telegram
        user_tier: User's tier (free or pro)
        
    Returns:
        Tuple of (is_valid, file_type, error_message)
    """
    file_extension = Path(file_name).suffix.lower()
    
    # DOCX format - supported for all tiers
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_extension == ".docx":
        return (True, "docx", "")
    
    # PDF format - only for premium users
    if mime_type == "application/pdf" or file_extension == ".pdf":
        if user_tier == "pro":
            return (True, "pdf", "")
        else:
            return (False, "", "📄 *PDF uploads are a Premium feature*\n\n"
                              "Free tier supports: ✅ DOCX files\n"
                              "Premium tier adds: ✅ PDF files\n\n"
                              "Please upload a .docx file, or upgrade to Premium to use PDF files!")
    
    # DOC format (old Word) - not supported yet
    if mime_type == "application/msword" or file_extension == ".doc":
        return (False, "", "❌ *Old Word format (.doc) not supported*\n\n"
                          "Please save your document as .docx format and upload again.\n\n"
                          "_In Microsoft Word: File → Save As → Choose 'Word Document (.docx)'_")
    
    # Unsupported format
    return (False, "", f"❌ *Unsupported file format*\n\n"
                       f"File: {file_name}\n"
                       f"Type: {mime_type}\n\n"
                       f"Supported formats:\n"
                       f"• Free tier: ✅ .docx\n"
                       f"• Premium: ✅ .docx, ✅ .pdf\n\n"
                       f"Please upload a valid resume file!")
