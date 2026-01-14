"""
CareerBuddy - Document Renderer (DOCX)
Document rendering service - Generate professional DOCX and PDF files
Author: Sir Dave
"""
from io import BytesIO
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from loguru import logger

from app.models import Job


def _clean_skills(skills_list):
    """Remove invalid skills (numbers, empty strings) from old data."""
    if not skills_list:
        return []
    return [s.strip() for s in skills_list if s and isinstance(s, str) and not s.strip().isdigit() and len(s.strip()) > 1]


def _add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: docx paragraph object
        url: URL to link to
        text: Display text for the link
    
    Returns:
        The hyperlink run
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    # Create the relationship
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink styling (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add hyperlink to paragraph
    paragraph._p.append(hyperlink)
    
    # Return a run object for further styling
    from docx.text.run import Run
    return Run(new_run, paragraph)


def render_resume(job: Job) -> bytes:
    """
    Generate professional DOCX resume with table-based layout.
    Supports multiple templates based on user selection.
    """
    answers = job.answers or {}
    selected_template = answers.get("template", "template_1")

    # Route to appropriate template renderer
    if selected_template == "template_2":
        return _render_template_2(answers)
    elif selected_template == "template_3":
        return _render_template_3(answers)
    else:
        # Default to template 1
        return _render_template_1(answers)


def _render_template_1(answers: dict) -> bytes:
    """
    Template 1: Classic Professional Layout
    - Centered header with NO icons
    - Calibri 12pt body, 14pt headings
    - Contact info separated by pipes (|)
    - Skills in 2 columns
    - Dates aligned far right
    - Medium spacing (12pt) between sections
    - 1.25" margins all round
    """
    doc = Document()

    # Set document margins - 0.5 inches all round (to match reference template)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Get data
    basics = answers.get('basics', {})
    summary = answers.get('summary', '')
    skills = _clean_skills(answers.get('skills', []))
    experiences = answers.get('experiences', [])
    education = answers.get('education', [])
    projects = answers.get('projects', [])
    profiles = answers.get('profiles', [])
    references = answers.get('references', [])

    # ==================== HEADER ====================
    name = basics.get('name', 'Your Name')
    title = basics.get('title', 'Professional Title')

    # Name (Large, Bold, Centered) - Increased font size
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.font.size = Pt(24)  # Increased from 20pt
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_para.space_after = Pt(0)  # Tighter spacing

    # Title (Centered, ALL CAPS) - Increased font size
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())  # ALL CAPS
    title_run.font.size = Pt(14)  # Increased from 12pt
    title_run.font.name = 'Calibri'
    title_para.space_after = Pt(0)  # Tighter spacing

    # Contact Info - NO ICONS, pipe separators - Calibri 12pt
    contact_parts = []
    if basics.get('location'):
        contact_parts.append(basics['location'])
    if basics.get('phone'):
        contact_parts.append(basics['phone'])
    if basics.get('email'):
        contact_parts.append(basics['email'])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.size = Pt(13)  # Increased from 12pt
        contact_run.font.name = 'Calibri'
        contact_para.space_after = Pt(0)  # Tighter spacing

    # Add horizontal line after header
    _add_horizontal_line(doc)
    doc.add_paragraph().space_after = Pt(10)  # Reduced spacing before first section

    # ==================== MAIN CONTENT TABLE ====================
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    # Calculate number of rows needed
    num_rows = 0
    if profiles: num_rows += 1
    if summary: num_rows += 1
    if experiences: num_rows += 1
    if education: num_rows += 1
    if projects: num_rows += 1
    if references: num_rows += 1
    if skills: num_rows += 1
    
    if num_rows == 0:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    # Remove all table borders (make them invisible)
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Set cell margins to zero to prevent extra spacing in PDF
    tblCellMar = OxmlElement('w:tblCellMar')
    for margin_name in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '0')
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    tblPr.append(tblCellMar)
    
    # Set column widths (0.5" margins = 7.5" available width)
    for row in table.rows:
        row.cells[0].width = Inches(1.2)  # Labels column
        row.cells[1].width = Inches(6.3)  # Content column (full width)
        
        # Set vertical alignment to top for both cells to reduce spacing
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    current_row = 0

    # ==================== PROFILES ====================
    if profiles:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Profiles')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_para = content_cell.paragraphs[0]
        content_para.paragraph_format.space_before = Pt(0)
        content_para.paragraph_format.space_after = Pt(0)  # Zero spacing - LibreOffice treats this as row spacing
        
        # Add clickable hyperlinks for profiles (NO icons)
        for idx, profile in enumerate(profiles):
            platform = profile.get('platform', 'Profile')
            url = profile.get('url', '')
            if platform and url:
                if idx > 0:
                    content_para.add_run(' | ')  # Separator
                
                # Add hyperlink
                hyperlink = _add_hyperlink(content_para, url, platform)
                hyperlink.font.size = Pt(12)
                hyperlink.font.name = 'Calibri'
        
        content_para.paragraph_format.space_after = Pt(16)  # Spacing between sections
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== SUMMARY ====================
    if summary:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Summary')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_para = content_cell.paragraphs[0]
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        content_para.paragraph_format.space_before = Pt(0)
        content_para.paragraph_format.space_after = Pt(16)  # Spacing between sections
        content_run = content_para.add_run(summary)
        content_run.font.size = Pt(12)  # Body size
        content_run.font.name = 'Calibri'
        
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== EXPERIENCE ====================
    if experiences:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Experience')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_cell._element.remove(content_cell.paragraphs[0]._element)
        
        for idx, exp in enumerate(experiences):
            # Company name and dates on same line
            exp_header = content_cell.add_paragraph()
            if idx == 0:
                exp_header.paragraph_format.space_before = Pt(0)  # No space before first item
            exp_header.paragraph_format.space_after = Pt(0)
            
            company = exp.get('company', 'Company Name')
            company_run = exp_header.add_run(company)
            company_run.font.bold = True
            company_run.font.size = Pt(12)
            company_run.font.name = 'Calibri'
            
            # Date aligned to far right - BOLD
            date_str = ''
            if exp.get('start'):
                date_str = exp['start']
                if exp.get('end'):
                    date_str += f" - {exp['end']}"
            
            if date_str:
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = exp_header.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)  # Right aligned with proper padding
                exp_header.add_run('\t')
                date_run = exp_header.add_run(date_str)
                date_run.font.bold = True  # Bold the date
                date_run.font.size = Pt(12)
                date_run.font.name = 'Calibri'
            
            # Job title and location on same line
            role_para = content_cell.add_paragraph()
            role_para.paragraph_format.space_after = Pt(0)
            
            role = exp.get('title', exp.get('role', 'Job Title'))
            role_run = role_para.add_run(role)
            role_run.font.size = Pt(12)
            role_run.font.name = 'Calibri'
            
            # Location on same line as role, right-aligned
            location = exp.get('city', exp.get('location', ''))
            if location:
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = role_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)  # Aligned with date
                role_para.add_run('\t')
                loc_run = role_para.add_run(location)
                loc_run.font.size = Pt(12)
                loc_run.font.name = 'Calibri'
            
            # Bullets with proper alignment
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                bullet_para = content_cell.add_paragraph()
                bullet_para.paragraph_format.space_after = Pt(0)
                bullet_para.paragraph_format.left_indent = Inches(0.2)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.15)
                bullet_run = bullet_para.add_run(f"• {bullet}")
                bullet_run.font.size = Pt(12)
                bullet_run.font.name = 'Calibri'
            
            # Spacing between experiences
            if idx < len(experiences) - 1:
                content_cell.add_paragraph().space_after = Pt(0)
        
        content_cell.paragraphs[-1].paragraph_format.space_after = Pt(16)  # Spacing between sections
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== EDUCATION ====================
    if education:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Education')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_cell._element.remove(content_cell.paragraphs[0]._element)
        
        for idx, edu in enumerate(education):
            # Institution (bold) on left, Date (bold) on far right
            edu_header = content_cell.add_paragraph()
            if idx == 0:
                edu_header.paragraph_format.space_before = Pt(0)  # No space before first item
            edu_header.paragraph_format.space_after = Pt(0)
            
            institution = edu.get('institution', 'Institution Name')
            inst_run = edu_header.add_run(institution)
            inst_run.font.bold = True  # Bold
            inst_run.font.size = Pt(12)
            inst_run.font.name = 'Calibri'
            
            years = edu.get('years', '')
            if years:
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = edu_header.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)  # Far right
                edu_header.add_run('\t')
                years_run = edu_header.add_run(years)
                years_run.font.bold = True  # Bold
                years_run.font.size = Pt(12)
                years_run.font.name = 'Calibri'
            
            # Course on left, Degree type on right
            degree_para = content_cell.add_paragraph()
            degree_para.paragraph_format.space_after = Pt(0)
            
            degree = edu.get('degree', '')
            degree_run = degree_para.add_run(degree)
            degree_run.font.size = Pt(12)
            degree_run.font.name = 'Calibri'
            
            degree_type = edu.get('degree_type', '')
            if degree_type:
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = degree_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT)  # Far right
                degree_para.add_run('\t')
                type_run = degree_para.add_run(degree_type)
                type_run.font.size = Pt(12)
                type_run.font.name = 'Calibri'
        
        content_cell.paragraphs[-1].paragraph_format.space_after = Pt(16)  # Spacing between sections
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== PROJECTS ====================
    if projects:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Projects')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_cell._element.remove(content_cell.paragraphs[0]._element)

        for idx, proj in enumerate(projects):
            proj_para = content_cell.add_paragraph()
            if idx == 0:
                proj_para.paragraph_format.space_before = Pt(0)  # No space before first item
            proj_para.paragraph_format.space_after = Pt(0)
            proj_para.paragraph_format.left_indent = Inches(0.2)
            proj_para.paragraph_format.first_line_indent = Inches(-0.15)
            proj_details = proj.get('details', '')
            if proj_details:
                proj_run = proj_para.add_run(f"• {proj_details}")
                proj_run.font.size = Pt(12)
                proj_run.font.name = 'Calibri'
        
        content_cell.paragraphs[-1].paragraph_format.space_after = Pt(16)  # Spacing between sections
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== REFERENCES ====================
    if references:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('References')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_cell._element.remove(content_cell.paragraphs[0]._element)
        
        for idx, ref in enumerate(references):
            name_para = content_cell.add_paragraph()
            if idx == 0:
                name_para.paragraph_format.space_before = Pt(0)  # No space before first item
            name_para.paragraph_format.space_after = Pt(0)
            name_run = name_para.add_run(ref.get('name', 'Reference Name'))
            name_run.font.bold = True
            name_run.font.size = Pt(12)
            name_run.font.name = 'Calibri'
            
            ref_title = ref.get('title', '')
            if ref_title:
                title_para = content_cell.add_paragraph(ref_title)
                title_para.paragraph_format.space_after = Pt(0)
                for run in title_para.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
            
            ref_org = ref.get('organization', '')
            if ref_org:
                org_para = content_cell.add_paragraph(ref_org)
                org_para.paragraph_format.space_after = Pt(0)
                for run in org_para.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
        
        content_cell.paragraphs[-1].paragraph_format.space_after = Pt(16)  # Spacing between sections
        _add_table_row_border(table.rows[current_row])
        current_row += 1

    # ==================== SKILLS ====================
    if skills:
        label_cell = table.rows[current_row].cells[0]
        content_cell = table.rows[current_row].cells[1]
        
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(0)  # Zero spacing to prevent PDF conversion gaps
        label_run = label_para.add_run('Skills')
        label_run.font.bold = True
        label_run.font.size = Pt(14)  # Heading size
        label_run.font.name = 'Calibri'
        
        content_cell._element.remove(content_cell.paragraphs[0]._element)
        
        # Split skills into 2 equal columns
        from docx.enum.text import WD_TAB_ALIGNMENT
        skills_per_col = (len(skills) + 1) // 2  # Round up
        col1 = skills[:skills_per_col]
        col2 = skills[skills_per_col:]
        
        # Create skill lines with 2 columns
        for i in range(max(len(col1), len(col2))):
            skill_line = content_cell.add_paragraph()
            if i == 0:
                skill_line.paragraph_format.space_before = Pt(0)  # No space before first item
            skill_line.paragraph_format.space_after = Pt(0)
            
            # Set up tab stop for 2 equal columns
            tab_stops = skill_line.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.75), WD_TAB_ALIGNMENT.LEFT)  # Column 2 start (adjusted for 0.5" margins)
            
            # Column 1 - BOLD
            if i < len(col1):
                skill1_run = skill_line.add_run(col1[i])
                skill1_run.font.bold = True  # Bold
                skill1_run.font.size = Pt(12)
                skill1_run.font.name = 'Calibri'
            
            skill_line.add_run('\t')
            
            # Column 2 - BOLD
            if i < len(col2):
                skill2_run = skill_line.add_run(col2[i])
                skill2_run.font.bold = True  # Bold
                skill2_run.font.size = Pt(12)
                skill2_run.font.name = 'Calibri'
        
        _add_table_row_border(table.rows[current_row])

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info("[renderer] Generated resume DOCX with template_1")
    return buffer.getvalue()


def _render_template_2(answers: dict) -> bytes:
    """
    Template 2: Modern Minimal Layout
    - Clean, contemporary design with accent colors
    - ALL CAPS section headings in dark blue
    - No icons, professional formatting
    - Calibri font, refined spacing
    """
    doc = Document()
    
    # Set margins to match template 1
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Get data
    basics = answers.get('basics', {})
    summary = answers.get('summary', '')
    skills = _clean_skills(answers.get('skills', []))[:6]  # Limit to 6 skills
    experiences = answers.get('experiences', [])
    education = answers.get('education', [])
    profiles = answers.get('profiles', [])
    certifications = answers.get('certifications', [])
    projects = answers.get('projects', [])
    
    # ==================== HEADER ====================
    # Name (centered, large, dark blue)
    name = basics.get('name', 'Your Name')
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(26)
    name_run.font.name = 'Calibri'
    name_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    name_para.paragraph_format.space_after = Pt(2)
    
    # Title (centered, gray)
    if basics.get('title'):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(basics['title'])
        title_run.font.size = Pt(13)
        title_run.font.name = 'Calibri'
        title_run.font.color.rgb = RGBColor(80, 80, 80)
        title_para.paragraph_format.space_after = Pt(2)
    
    # Contact info (centered, no icons)
    contact_parts = []
    if basics.get('location'):
        contact_parts.append(basics['location'])
    if basics.get('phone'):
        contact_parts.append(basics['phone'])
    if basics.get('email'):
        contact_parts.append(basics['email'])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.size = Pt(11)
        contact_run.font.name = 'Calibri'
        contact_run.font.color.rgb = RGBColor(60, 60, 60)
        contact_para.paragraph_format.space_after = Pt(16)
    
    # Horizontal line separator
    _add_horizontal_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ==================== PROFILES ====================
    if profiles:
        profiles_heading = doc.add_paragraph()
        profiles_heading_run = profiles_heading.add_run('PROFILES')
        profiles_heading_run.bold = True
        profiles_heading_run.font.size = Pt(13)
        profiles_heading_run.font.name = 'Calibri'
        profiles_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        profiles_heading.paragraph_format.space_after = Pt(6)
        
        profiles_para = doc.add_paragraph()
        for idx, profile in enumerate(profiles):
            platform = profile.get('platform', 'Profile')
            url = profile.get('url', '')
            if platform and url:
                if idx > 0:
                    profiles_para.add_run(' | ')
                _add_hyperlink(profiles_para, url, platform)
        profiles_para.paragraph_format.space_after = Pt(14)
    
    # ==================== SUMMARY ====================
    if summary:
        summary_heading = doc.add_paragraph()
        summary_heading_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
        summary_heading_run.bold = True
        summary_heading_run.font.size = Pt(13)
        summary_heading_run.font.name = 'Calibri'
        summary_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        summary_heading.paragraph_format.space_after = Pt(6)
        
        summary_para = doc.add_paragraph(summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_para.paragraph_format.space_after = Pt(14)
        for run in summary_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    
    # ==================== SKILLS ====================
    if skills:
        skills_heading = doc.add_paragraph()
        skills_heading_run = skills_heading.add_run('KEY SKILLS')
        skills_heading_run.bold = True
        skills_heading_run.font.size = Pt(13)
        skills_heading_run.font.name = 'Calibri'
        skills_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        skills_heading.paragraph_format.space_after = Pt(6)
        
        skills_text = ' • '.join(skills)
        skills_para = doc.add_paragraph(skills_text)
        skills_para.paragraph_format.space_after = Pt(14)
        for run in skills_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    
    # ==================== EXPERIENCE ====================
    if experiences:
        exp_heading = doc.add_paragraph()
        exp_heading_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
        exp_heading_run.bold = True
        exp_heading_run.font.size = Pt(13)
        exp_heading_run.font.name = 'Calibri'
        exp_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        exp_heading.paragraph_format.space_after = Pt(6)
        
        for idx, exp in enumerate(experiences):
            # Company (bold) and dates
            comp_para = doc.add_paragraph()
            comp_run = comp_para.add_run(exp.get('company', 'Company'))
            comp_run.bold = True
            comp_run.font.size = Pt(12)
            comp_run.font.name = 'Calibri'
            
            dates = f"{exp.get('start', '')} - {exp.get('end', '')}"
            if dates.strip() != ' - ':
                comp_para.add_run(f"  |  {dates}")
            comp_para.paragraph_format.space_after = Pt(2)
            
            # Title and location
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', exp.get('role', 'Position')))
            title_run.italic = True
            title_run.font.size = Pt(11)
            title_run.font.name = 'Calibri'
            
            location = exp.get('city', exp.get('location', ''))
            if location:
                title_para.add_run(f" – {location}")
            title_para.paragraph_format.space_after = Pt(4)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                bullet_para = doc.add_paragraph()
                bullet_para.paragraph_format.left_indent = Inches(0.2)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.15)
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_para.add_run(f"• {bullet}")
                bullet_run.font.size = Pt(11)
                bullet_run.font.name = 'Calibri'
            
            # Spacer between jobs
            if idx < len(experiences) - 1:
                doc.add_paragraph().paragraph_format.space_after = Pt(8)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ==================== EDUCATION ====================
    if education:
        edu_heading = doc.add_paragraph()
        edu_heading_run = edu_heading.add_run('EDUCATION')
        edu_heading_run.bold = True
        edu_heading_run.font.size = Pt(13)
        edu_heading_run.font.name = 'Calibri'
        edu_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        edu_heading.paragraph_format.space_after = Pt(6)
        
        for edu in education:
            institution = edu.get('institution', 'Institution')
            degree = edu.get('degree', '')
            years = edu.get('years', '')
            
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(4)
            
            # Institution and year
            inst_run = edu_para.add_run(f"{institution}")
            inst_run.bold = True
            inst_run.font.size = Pt(11)
            inst_run.font.name = 'Calibri'
            
            if years:
                edu_para.add_run(f"  |  {years}")
            
            # Degree
            if degree:
                degree_para = doc.add_paragraph(degree)
                degree_para.paragraph_format.space_after = Pt(6)
                for run in degree_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
    
    # ==================== CERTIFICATIONS ====================
    if certifications:
        cert_heading = doc.add_paragraph()
        cert_heading_run = cert_heading.add_run('CERTIFICATIONS')
        cert_heading_run.bold = True
        cert_heading_run.font.size = Pt(13)
        cert_heading_run.font.name = 'Calibri'
        cert_heading_run.font.color.rgb = RGBColor(0, 51, 102)
        cert_heading.paragraph_format.space_after = Pt(6)
        
        for cert in certifications:
            cert_name = cert.get('name', '')
            cert_body = cert.get('issuing_body', '')
            cert_year = cert.get('year', '')
            
            cert_text = f"{cert_name}, {cert_body}"
            if cert_year:
                cert_text += f", {cert_year}"
            
            cert_para = doc.add_paragraph(f"• {cert_text}")
            cert_para.paragraph_format.space_after = Pt(4)
            for run in cert_para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("[renderer] Generated resume DOCX with template_2")
    return buffer.getvalue()


def _render_template_3(answers: dict) -> bytes:
    """
    Template 3: Executive Bold Layout
    - Strong visual hierarchy with bold black section headers
    - Larger fonts and generous spacing for executive presence
    - Authoritative, commanding appearance
    - Arial Black/Arial for bold impact
    """
    doc = Document()
    
    # Set generous margins for executive presence
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Get data
    basics = answers.get('basics', {})
    summary = answers.get('summary', '')
    skills = _clean_skills(answers.get('skills', []))[:6]  # Limit to 6 skills
    experiences = answers.get('experiences', [])
    education = answers.get('education', [])
    profiles = answers.get('profiles', [])
    certifications = answers.get('certifications', [])
    projects = answers.get('projects', [])
    
    # ==================== HEADER ====================
    # Name (bold, very large, ALL CAPS)
    name = basics.get('name', 'Your Name')
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_para.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(28)
    name_run.font.name = 'Arial'
    name_para.paragraph_format.space_after = Pt(4)
    
    # Title (larger, bold)
    if basics.get('title'):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(basics['title'].upper())
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Arial'
        title_run.font.color.rgb = RGBColor(60, 60, 60)
        title_para.paragraph_format.space_after = Pt(4)
    
    # Contact info (no icons, pipe separated)
    contact_parts = []
    if basics.get('location'):
        contact_parts.append(basics['location'])
    if basics.get('phone'):
        contact_parts.append(basics['phone'])
    if basics.get('email'):
        contact_parts.append(basics['email'])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.size = Pt(11)
        contact_run.font.name = 'Arial'
        contact_run.font.color.rgb = RGBColor(80, 80, 80)
        contact_para.paragraph_format.space_after = Pt(18)
    
    # Thick horizontal line separator
    _add_horizontal_line(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    
    # ==================== PROFILES ====================
    if profiles:
        profiles_heading = doc.add_paragraph()
        profiles_heading_run = profiles_heading.add_run('PROFESSIONAL PROFILES')
        profiles_heading_run.bold = True
        profiles_heading_run.font.size = Pt(14)
        profiles_heading_run.font.name = 'Arial'
        profiles_heading.paragraph_format.space_after = Pt(8)
        
        profiles_para = doc.add_paragraph()
        for idx, profile in enumerate(profiles):
            platform = profile.get('platform', 'Profile')
            url = profile.get('url', '')
            if platform and url:
                if idx > 0:
                    profiles_para.add_run('  |  ')
                _add_hyperlink(profiles_para, url, platform)
        profiles_para.paragraph_format.space_after = Pt(16)
    
    # ==================== SUMMARY ====================
    if summary:
        summary_heading = doc.add_paragraph()
        summary_heading_run = summary_heading.add_run('EXECUTIVE SUMMARY')
        summary_heading_run.bold = True
        summary_heading_run.font.size = Pt(14)
        summary_heading_run.font.name = 'Arial'
        summary_heading.paragraph_format.space_after = Pt(8)
        
        summary_para = doc.add_paragraph(summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_para.paragraph_format.space_after = Pt(16)
        for run in summary_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
    
    # ==================== CORE COMPETENCIES ====================
    if skills:
        skills_heading = doc.add_paragraph()
        skills_heading_run = skills_heading.add_run('CORE COMPETENCIES')
        skills_heading_run.bold = True
        skills_heading_run.font.size = Pt(14)
        skills_heading_run.font.name = 'Arial'
        skills_heading.paragraph_format.space_after = Pt(8)
        
        skills_text = '  •  '.join(skills)
        skills_para = doc.add_paragraph(skills_text)
        skills_para.paragraph_format.space_after = Pt(16)
        for run in skills_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.bold = True
    
    # ==================== PROFESSIONAL EXPERIENCE ====================
    if experiences:
        exp_heading = doc.add_paragraph()
        exp_heading_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
        exp_heading_run.bold = True
        exp_heading_run.font.size = Pt(14)
        exp_heading_run.font.name = 'Arial'
        exp_heading.paragraph_format.space_after = Pt(8)
        
        for idx, exp in enumerate(experiences):
            # Company (bold, ALL CAPS, larger)
            comp_para = doc.add_paragraph()
            comp_run = comp_para.add_run(exp.get('company', 'Company').upper())
            comp_run.bold = True
            comp_run.font.size = Pt(12)
            comp_run.font.name = 'Arial'
            comp_para.paragraph_format.space_before = Pt(8) if idx > 0 else Pt(0)
            comp_para.paragraph_format.space_after = Pt(3)
            
            # Title and dates
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get('title', exp.get('role', 'Position')))
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = 'Arial'
            
            dates = f"{exp.get('start', '')} - {exp.get('end', '')}"
            if dates.strip() != ' - ':
                title_para.add_run(f"  |  {dates}")
            title_para.paragraph_format.space_after = Pt(2)
            
            # Location
            location = exp.get('city', exp.get('location', ''))
            if location:
                loc_para = doc.add_paragraph(location)
                loc_para.paragraph_format.space_after = Pt(5)
                for run in loc_para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.italic = True
                    run.font.color.rgb = RGBColor(80, 80, 80)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                bullet_para = doc.add_paragraph()
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.2)
                bullet_para.paragraph_format.space_after = Pt(4)
                bullet_run = bullet_para.add_run(f"▪ {bullet}")
                bullet_run.font.size = Pt(11)
                bullet_run.font.name = 'Arial'
        
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ==================== EDUCATION ====================
    if education:
        edu_heading = doc.add_paragraph()
        edu_heading_run = edu_heading.add_run('EDUCATION')
        edu_heading_run.bold = True
        edu_heading_run.font.size = Pt(14)
        edu_heading_run.font.name = 'Arial'
        edu_heading.paragraph_format.space_after = Pt(8)
        
        for edu in education:
            institution = edu.get('institution', 'Institution')
            degree = edu.get('degree', '')
            years = edu.get('years', '')
            
            # Institution and year
            edu_para = doc.add_paragraph()
            inst_run = edu_para.add_run(f"{institution}")
            inst_run.bold = True
            inst_run.font.size = Pt(11)
            inst_run.font.name = 'Arial'
            
            if years:
                edu_para.add_run(f"  |  {years}")
            edu_para.paragraph_format.space_after = Pt(3)
            
            # Degree
            if degree:
                degree_para = doc.add_paragraph(degree)
                degree_para.paragraph_format.space_after = Pt(8)
                for run in degree_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    # ==================== CERTIFICATIONS ====================
    if certifications:
        cert_heading = doc.add_paragraph()
        cert_heading_run = cert_heading.add_run('CERTIFICATIONS & CREDENTIALS')
        cert_heading_run.bold = True
        cert_heading_run.font.size = Pt(14)
        cert_heading_run.font.name = 'Arial'
        cert_heading.paragraph_format.space_after = Pt(8)
        
        for cert in certifications:
            cert_name = cert.get('name', '')
            cert_body = cert.get('issuing_body', '')
            cert_year = cert.get('year', '')
            
            cert_text = f"{cert_name}"
            if cert_body:
                cert_text += f", {cert_body}"
            if cert_year:
                cert_text += f", {cert_year}"
            
            cert_para = doc.add_paragraph(f"▪ {cert_text}")
            cert_para.paragraph_format.left_indent = Inches(0.25)
            cert_para.paragraph_format.space_after = Pt(5)
            for run in cert_para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    # ==================== PROJECTS ====================
    if projects:
        proj_heading = doc.add_paragraph()
        proj_heading_run = proj_heading.add_run('KEY PROJECTS')
        proj_heading_run.bold = True
        proj_heading_run.font.size = Pt(14)
        proj_heading_run.font.name = 'Arial'
        proj_heading.paragraph_format.space_after = Pt(8)
        
        for proj in projects:
            proj_name = proj.get('name', '')
            proj_details = proj.get('details', '')
            
            if proj_name and proj_details:
                proj_para = doc.add_paragraph()
                proj_name_run = proj_para.add_run(f"{proj_name}: ")
                proj_name_run.bold = True
                proj_para.add_run(proj_details)
                proj_para.paragraph_format.space_after = Pt(6)
                for run in proj_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("[renderer] Generated resume DOCX with template_3")
    return buffer.getvalue()


def render_cv(job: Job) -> bytes:
    """
    Generate CV with professional layout.
    CV now uses the EXACT SAME layout as Resume for consistency.
    The only difference is the document type label - layout is identical.
    """
    answers = job.answers or {}
    selected_template = answers.get("template", "template_1")
    
    # Use the same template renderers as Resume
    if selected_template == "template_2":
        return _render_template_2(answers)
    elif selected_template == "template_3":
        return _render_template_3(answers)
    else:
        # Use the same Template 1 as Resume (perfected layout)
        return _render_template_1(answers)


# NOTE: Old _render_cv_template_1 function removed - CV now uses same layout as Resume



def render_cover_letter(job: Job) -> bytes:
    """
    Generate professional cover letter DOCX following HR industry standards.
    """
    answers = job.answers or {}
    basics = answers.get("basics", {})
    
    # Extract all data
    name = basics.get("name", "Your Name")
    email = basics.get("email", "")
    phone = basics.get("phone", "")
    location = basics.get("location", "")
    linkedin = basics.get("linkedin", "")
    
    role = answers.get("cover_role", "Position Title")
    company = answers.get("cover_company", "Company Name")
    years_exp = answers.get("years_experience", "X years")
    industries = answers.get("industries", "relevant industries")
    interest_reason = answers.get("interest_reason", "this exciting opportunity")
    current_title = answers.get("current_title", "Your Current Title")
    current_employer = answers.get("current_employer", "Current Employer")
    achievement_1 = answers.get("achievement_1", "")
    achievement_2 = answers.get("achievement_2", "")
    key_skills = answers.get("cover_key_skills", [])
    company_goal = answers.get("company_goal", "the organization's goals")

    doc = Document()

    # Set margins - professional business letter format
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ==================== HEADER (Contact Info) ====================
    # Name
    name_para = doc.add_paragraph(name)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(2)
    for run in name_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
    
    # Contact line: LinkedIn | Phone | Email
    contact_parts = []
    if linkedin:
        contact_parts.append(linkedin)
    if phone:
        contact_parts.append(phone)
    if email:
        contact_parts.append(email)
    
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_para.paragraph_format.space_after = Pt(20)
        for run in contact_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    # ==================== DATE ====================
    from datetime import datetime
    date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_para.paragraph_format.space_after = Pt(20)
    for run in date_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # ==================== GREETING ====================
    greeting_para = doc.add_paragraph("Dear Hiring Manager,")
    greeting_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    greeting_para.paragraph_format.space_after = Pt(12)
    for run in greeting_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # ==================== OPENING PARAGRAPH ====================
    opening = (
        f"I am writing to express my interest in the {role} role at {company}. "
        f"With {years_exp} of progressive experience across {industries}, I deliver measurable "
        f"improvements in key areas of responsibility. I am particularly excited about this "
        f"opportunity because {interest_reason}."
    )
    opening_para = doc.add_paragraph(opening)
    opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    opening_para.paragraph_format.space_after = Pt(12)
    for run in opening_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # ==================== BODY PARAGRAPH (Experience & Achievements) ====================
    body_text = (
        f"In my current/most recent role as {current_title} at {current_employer}, "
        f"{achievement_1}"
    )
    
    if achievement_2:
        body_text += f" I also {achievement_2}"
    
    body_para = doc.add_paragraph(body_text)
    body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_para.paragraph_format.space_after = Pt(12)
    for run in body_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # ==================== SKILLS PARAGRAPH ====================
    if key_skills:
        skills_text = ", ".join(key_skills[:-1])
        if len(key_skills) > 1:
            skills_text += f", and {key_skills[-1]}"
        else:
            skills_text = key_skills[0] if key_skills else ""
        
        skills_para_text = (
            f"I bring a pragmatic combination of strategic leadership and hands-on operational "
            f"execution: {skills_text}. I would welcome the chance to discuss how my experience "
            f"in these areas can support {company}'s goals for {company_goal}."
        )
        
        skills_para = doc.add_paragraph(skills_para_text)
        skills_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        skills_para.paragraph_format.space_after = Pt(12)
        for run in skills_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    # ==================== CLOSING PARAGRAPH ====================
    closing_text = (
        "Thank you for considering my application. I have attached my resume and I am available "
        "for a conversation at your convenience. I look forward to the possibility of contributing "
        "to your team."
    )
    closing_para = doc.add_paragraph(closing_text)
    closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing_para.paragraph_format.space_after = Pt(20)
    for run in closing_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # ==================== SIGNATURE ====================
    sincerely_para = doc.add_paragraph("Sincerely,")
    sincerely_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sincerely_para.paragraph_format.space_after = Pt(20)
    for run in sincerely_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    signature_para = doc.add_paragraph(name)
    signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    signature_para.paragraph_format.space_after = Pt(2)
    for run in signature_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    # Contact line in signature
    if contact_parts:
        sig_contact_para = doc.add_paragraph(" | ".join(contact_parts))
        sig_contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in sig_contact_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.info(f"[renderer] Generated professional cover letter DOCX for job.id={job.id}")
    return buffer.getvalue()


def render_revamp(job: Job) -> bytes:
    """
    Generate a DOCX file from AI-revamped resume content (free-form text).
    """
    answers = job.answers or {}
    improved_text = answers.get("revamped_content") or answers.get("original_content") or ""

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    _add_section_heading(doc, "IMPROVED RESUME CONTENT")

    if not improved_text:
        para = doc.add_paragraph("No content provided.")
        _set_body_font(para)
    else:
        for line in improved_text.splitlines():
            if not line.strip():
                doc.add_paragraph()
                continue
            para = doc.add_paragraph(line.strip())
            _set_body_font(para)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    logger.info(f"[renderer] Generated revamped DOCX for job.id={job.id}")
    return buffer.getvalue()


def _add_section_heading(doc: Document, text: str):
    """Add a section heading with consistent styling"""
    heading = doc.add_paragraph()
    heading_run = heading.add_run(text)
    heading_run.font.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.name = 'Arial'
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add bottom border
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Border width
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_cv_section_heading(doc: Document, text: str):
    """Add a CV section heading with horizontal line separator"""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(4)
    heading_run = heading.add_run(text)
    heading_run.font.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.name = 'Arial'
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add bottom border (horizontal line)
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # Thinner border
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_body_font(paragraph):
    """Set consistent body text formatting"""
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)


def _add_horizontal_line(doc: Document):
    """Add a horizontal line separator"""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(0)


def _add_table_row_border(row):
    """Add bottom border to a table row"""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    for cell in row.cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)
