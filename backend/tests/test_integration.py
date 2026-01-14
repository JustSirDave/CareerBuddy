"""
Integration tests for complete workflows
Tests end-to-end scenarios across multiple components
"""
import pytest
from unittest.mock import patch, Mock
from app.services.router import handle_inbound
from app.services import renderer, payments
from app.models import User, Job, Message
from io import BytesIO
from docx import Document


@pytest.mark.asyncio
@pytest.mark.integration
class TestCompleteResumeFlow:
    """Test complete resume creation flow from start to finish"""

    async def test_full_resume_creation(self, db_session):
        """Test complete resume creation from greeting to document"""
        telegram_id = "integration_test_001"
        
        # Step 1: Welcome message
        response = await handle_inbound(db_session, telegram_id, "/start")
        assert "Welcome" in response or "welcome" in response.lower()
        
        # Verify user created
        user = db_session.query(User).filter(
            User.telegram_user_id == telegram_id
        ).first()
        assert user is not None
        
        # Step 2: Start resume
        response = await handle_inbound(db_session, telegram_id, "resume")
        assert "name" in response.lower() or "email" in response.lower()
        
        # Step 3: Provide basics
        response = await handle_inbound(
            db_session,
            telegram_id,
            "John Integration, john@integration.com, +1234567890, Test City"
        )
        assert response is not None
        
        # Step 4: Target role
        response = await handle_inbound(db_session, telegram_id, "Backend Engineer")
        assert response is not None
        
        # Step 5: Experience
        response = await handle_inbound(
            db_session,
            telegram_id,
            "Senior Developer, TechCorp, NYC, Jan 2020, Present"
        )
        assert "bullet" in response.lower() or "achievement" in response.lower()
        
        # Step 6: Bullets
        response = await handle_inbound(
            db_session,
            telegram_id,
            "Built scalable API serving 1M requests/day"
        )
        response = await handle_inbound(db_session, telegram_id, "done")
        
        # Step 7: Education
        response = await handle_inbound(
            db_session,
            telegram_id,
            "B.Sc. Computer Science, MIT, 2020"
        )
        response = await handle_inbound(db_session, telegram_id, "skip")
        
        # Verify job progression
        job = db_session.query(Job).filter(
            Job.user_id == user.id,
            Job.type == "resume"
        ).first()
        
        assert job is not None
        assert job.answers.get("basics") is not None
        assert job.answers["basics"]["name"] == "John Integration"

    async def test_resume_with_ai_skills(self, db_session):
        """Test resume flow with AI skills generation"""
        telegram_id = "integration_test_002"
        
        with patch('app.services.ai.generate_skills') as mock_skills:
            mock_skills.return_value = ["Python", "FastAPI", "PostgreSQL", "Docker", "AWS"]
            
            # Create user and job
            response = await handle_inbound(db_session, telegram_id, "/start")
            response = await handle_inbound(db_session, telegram_id, "resume")
            
            # Provide minimal data to reach skills step
            response = await handle_inbound(
                db_session,
                telegram_id,
                "Test User, test@test.com, +1234567890, City"
            )
            response = await handle_inbound(db_session, telegram_id, "Software Engineer")
            
            # AI skills should be generated
            # User selects skills
            user = db_session.query(User).filter(
                User.telegram_user_id == telegram_id
            ).first()
            
            job = db_session.query(Job).filter(
                Job.user_id == user.id
            ).first()
            
            # Simulate skill selection if we reach that step
            if job and "_step" in job.answers:
                assert job.answers.get("target_role") == "Software Engineer"


@pytest.mark.asyncio
@pytest.mark.integration
class TestCompleteCVFlow:
    """Test complete CV creation flow"""

    async def test_cv_creation_flow(self, db_session):
        """Test CV creation from start to finish"""
        telegram_id = "integration_cv_001"
        
        response = await handle_inbound(db_session, telegram_id, "/start")
        response = await handle_inbound(db_session, telegram_id, "cv")
        
        # Provide basics
        response = await handle_inbound(
            db_session,
            telegram_id,
            "Jane CV, jane@cv.com, +9876543210, London"
        )
        
        # Verify CV job created
        user = db_session.query(User).filter(
            User.telegram_user_id == telegram_id
        ).first()
        
        job = db_session.query(Job).filter(
            Job.user_id == user.id,
            Job.type == "cv"
        ).first()
        
        assert job is not None
        assert job.type == "cv"


@pytest.mark.asyncio
@pytest.mark.integration
class TestCompleteCoverLetterFlow:
    """Test complete cover letter creation flow"""

    async def test_cover_letter_creation(self, db_session):
        """Test cover letter from start to finish"""
        telegram_id = "integration_cover_001"
        
        response = await handle_inbound(db_session, telegram_id, "/start")
        response = await handle_inbound(db_session, telegram_id, "cover letter")
        
        # Provide basics
        response = await handle_inbound(
            db_session,
            telegram_id,
            "Alex Cover, alex@cover.com, +1111111111, Boston"
        )
        
        # Role and company
        response = await handle_inbound(
            db_session,
            telegram_id,
            "Senior Manager, Google"
        )
        
        # Verify cover letter job
        user = db_session.query(User).filter(
            User.telegram_user_id == telegram_id
        ).first()
        
        job = db_session.query(Job).filter(
            Job.user_id == user.id,
            Job.type == "cover"
        ).first()
        
        assert job is not None
        assert job.answers.get("cover_company") == "Google"


@pytest.mark.integration
class TestDocumentGenerationIntegration:
    """Test document generation with real data"""

    def test_generate_all_document_types(self, db_session, test_user, sample_resume_data, sample_cv_data, sample_cover_letter_data):
        """Test generating all document types"""
        # Resume
        resume_job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        resume_bytes = renderer.render_resume(resume_job)
        assert len(resume_bytes) > 0
        
        # CV
        cv_job = Job(
            user_id=test_user.id,
            type="cv",
            answers=sample_cv_data
        )
        cv_bytes = renderer.render_cv(cv_job)
        assert len(cv_bytes) > 0
        
        # Cover Letter
        cover_job = Job(
            user_id=test_user.id,
            type="cover",
            answers=sample_cover_letter_data
        )
        cover_bytes = renderer.render_cover_letter(cover_job)
        assert len(cover_bytes) > 0

    def test_document_content_validity(self, db_session, test_user, sample_resume_data):
        """Test generated document has valid content"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        
        # Extract all text
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Verify key content is present
        assert sample_resume_data["basics"]["name"] in full_text
        assert sample_resume_data["experiences"][0]["company"] in full_text
        assert any(skill in full_text for skill in sample_resume_data["skills"])


@pytest.mark.asyncio
@pytest.mark.integration
class TestPaymentFlow:
    """Test payment integration"""

    async def test_free_to_pro_upgrade(self, db_session, test_user):
        """Test upgrading from free to pro"""
        assert test_user.tier == "free"
        
        # Use payment bypass
        response = await handle_inbound(
            db_session,
            test_user.telegram_user_id,
            "payment made"
        )
        
        db_session.refresh(test_user)
        assert test_user.tier == "pro"
        assert "upgraded" in response.lower() or "pro" in response.lower()

    def test_pro_user_unlimited_generation(self, db_session, pro_user):
        """Test pro user has unlimited generations"""
        # Set high generation count
        pro_user.generation_count = 100
        db_session.commit()
        
        can_gen, reason = payments.can_generate(db_session, pro_user, "Any Role")
        assert can_gen is True

    def test_free_user_hits_limit(self, db_session, test_user):
        """Test free user hits generation limit"""
        # Max out generations
        test_user.generation_count = payments.MAX_FREE_GENERATIONS
        db_session.commit()
        
        can_gen, reason = payments.can_generate(db_session, test_user, "Backend Engineer")
        assert can_gen is False


@pytest.mark.integration
class TestDatabaseIntegration:
    """Test database operations integration"""

    def test_user_job_cascade(self, db_session, test_user):
        """Test cascading deletes work correctly"""
        # Create multiple jobs
        jobs = [
            Job(user_id=test_user.id, type="resume"),
            Job(user_id=test_user.id, type="cv"),
            Job(user_id=test_user.id, type="cover")
        ]
        for job in jobs:
            db_session.add(job)
        db_session.commit()
        
        job_ids = [j.id for j in jobs]
        
        # Delete user
        db_session.delete(test_user)
        db_session.commit()
        
        # All jobs should be deleted
        for job_id in job_ids:
            job = db_session.query(Job).filter(Job.id == job_id).first()
            assert job is None

    def test_message_tracking(self, db_session, test_job):
        """Test message history tracking"""
        messages = [
            Message(job_id=test_job.id, role="user", content="Hello"),
            Message(job_id=test_job.id, role="assistant", content="Hi"),
            Message(job_id=test_job.id, role="user", content="Resume please"),
        ]
        for msg in messages:
            db_session.add(msg)
        db_session.commit()
        
        db_session.refresh(test_job)
        assert len(test_job.messages) == 3


@pytest.mark.asyncio
@pytest.mark.integration
class TestErrorRecovery:
    """Test error recovery and resilience"""

    async def test_resume_flow_with_reset(self, db_session):
        """Test resetting and restarting flow"""
        telegram_id = "error_recovery_001"
        
        # Start flow
        await handle_inbound(db_session, telegram_id, "/start")
        await handle_inbound(db_session, telegram_id, "resume")
        await handle_inbound(db_session, telegram_id, "Test, test@test.com, +1234567890, City")
        
        # Reset
        response = await handle_inbound(db_session, telegram_id, "/reset")
        assert "reset" in response.lower() or "fresh" in response.lower()
        
        # Start again
        response = await handle_inbound(db_session, telegram_id, "resume")
        assert response is not None

    async def test_invalid_input_recovery(self, db_session, test_user, test_job):
        """Test recovery from invalid inputs"""
        # Provide invalid format
        test_job.answers = {"_step": "basics"}
        db_session.commit()
        
        from app.services.router import handle_resume
        response = await handle_resume(db_session, test_job, "invalid", "free")
        
        # Should re-prompt
        assert "format" in response.lower() or "example" in response.lower()
        
        # Should still be on same step
        db_session.refresh(test_job)
        assert test_job.answers["_step"] == "basics"


@pytest.mark.integration
class TestConcurrency:
    """Test concurrent operations"""

    @pytest.mark.asyncio
    async def test_multiple_users_simultaneously(self, db_session):
        """Test handling multiple users at once"""
        user_ids = [f"concurrent_{i}" for i in range(5)]
        
        # All users start flows
        for user_id in user_ids:
            response = await handle_inbound(db_session, user_id, "/start")
            assert response is not None
        
        # Verify all users created
        for user_id in user_ids:
            user = db_session.query(User).filter(
                User.telegram_user_id == user_id
            ).first()
            assert user is not None

    def test_concurrent_job_creation(self, db_session, test_user):
        """Test creating multiple jobs for same user"""
        jobs = [
            Job(user_id=test_user.id, type="resume", status="collecting"),
            Job(user_id=test_user.id, type="cv", status="collecting"),
            Job(user_id=test_user.id, type="cover", status="collecting")
        ]
        
        for job in jobs:
            db_session.add(job)
        db_session.commit()
        
        # All should be created
        user_jobs = db_session.query(Job).filter(
            Job.user_id == test_user.id
        ).all()
        
        assert len(user_jobs) >= 3


@pytest.mark.integration
class TestTemplateRendering:
    """Test all templates render correctly"""

    def test_all_templates_for_resume(self, db_session, test_user, sample_resume_data):
        """Test all three templates render resumes"""
        for template in ["template_1", "template_2", "template_3"]:
            data = sample_resume_data.copy()
            data["template"] = template
            
            job = Job(user_id=test_user.id, type="resume", answers=data)
            doc_bytes = renderer.render_resume(job)
            
            assert len(doc_bytes) > 0
            
            # Verify it's valid DOCX
            doc = Document(BytesIO(doc_bytes))
            assert len(doc.paragraphs) > 0


@pytest.mark.integration
class TestSystemLimits:
    """Test system handles limits and constraints"""

    def test_very_long_resume(self, db_session, test_user):
        """Test rendering resume with many experiences"""
        data = {
            "basics": {"name": "Experienced Pro", "email": "pro@test.com"},
            "summary": "Very experienced professional",
            "skills": ["Skill1", "Skill2", "Skill3"],
            "experiences": [
                {
                    "company": f"Company {i}",
                    "role": f"Role {i}",
                    "start": f"Jan {2010+i}",
                    "end": "Present" if i == 9 else f"Dec {2010+i}",
                    "location": "City",
                    "bullets": [f"Achievement {i}.1", f"Achievement {i}.2"]
                }
                for i in range(10)
            ],
            "education": [{"details": "PhD in Computer Science, Harvard, 2010"}],
            "projects": [{"details": f"Project {i}"} for i in range(5)]
        }
        
        job = Job(user_id=test_user.id, type="resume", answers=data)
        doc_bytes = renderer.render_resume(job)
        
        # Should handle large documents
        assert len(doc_bytes) > 0
        assert len(doc_bytes) < 5_000_000  # Less than 5MB

    def test_generation_limit_enforcement(self, db_session, test_user):
        """Test generation limits are enforced"""
        test_user.generation_count = 0
        db_session.commit()
        
        # Track generations
        for i in range(payments.MAX_FREE_GENERATIONS):
            payments.update_generation_count(db_session, test_user, f"Role{i}")
        
        db_session.refresh(test_user)
        assert test_user.generation_count >= payments.MAX_FREE_GENERATIONS
        
        # Should not be able to generate more
        can_gen, reason = payments.can_generate(db_session, test_user, "NewRole")
        assert can_gen is False
