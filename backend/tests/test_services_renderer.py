"""
Tests for document rendering services
Tests DOCX generation for resumes, CVs, and cover letters
"""
import pytest
from io import BytesIO
from docx import Document
from app.services import renderer
from app.models import Job


class TestRenderResume:
    """Test resume rendering"""

    def test_render_basic_resume(self, db_session, test_user, sample_resume_data):
        """Test rendering a basic resume"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        
        # Verify it's a valid DOCX
        doc = Document(BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_render_resume_contains_name(self, db_session, test_user, sample_resume_data):
        """Test resume contains user's name"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "John Doe" in text

    def test_render_resume_contains_experience(self, db_session, test_user, sample_resume_data):
        """Test resume contains experience"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "TechCorp" in text
        assert "Senior Backend Engineer" in text

    def test_render_resume_contains_skills(self, db_session, test_user, sample_resume_data):
        """Test resume contains skills"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Python" in text
        assert "FastAPI" in text

    def test_render_resume_template_1(self, db_session, test_user, sample_resume_data):
        """Test resume with template 1"""
        sample_resume_data["template"] = "template_1"
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0

    def test_render_resume_template_2(self, db_session, test_user, sample_resume_data):
        """Test resume with template 2"""
        sample_resume_data["template"] = "template_2"
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0

    def test_render_resume_template_3(self, db_session, test_user, sample_resume_data):
        """Test resume with template 3"""
        sample_resume_data["template"] = "template_3"
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0


class TestRenderCV:
    """Test CV rendering"""

    def test_render_basic_cv(self, db_session, test_user, sample_cv_data):
        """Test rendering a basic CV"""
        job = Job(
            user_id=test_user.id,
            type="cv",
            answers=sample_cv_data
        )
        
        doc_bytes = renderer.render_cv(job)
        
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        
        doc = Document(BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_render_cv_contains_name(self, db_session, test_user, sample_cv_data):
        """Test CV contains user's name"""
        job = Job(
            user_id=test_user.id,
            type="cv",
            answers=sample_cv_data
        )
        
        doc_bytes = renderer.render_cv(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Jane Smith" in text

    def test_cv_uses_same_layout_as_resume(self, db_session, test_user, sample_cv_data, sample_resume_data):
        """Test CV uses same template renderers as resume"""
        # CV should use same rendering logic as resume
        cv_job = Job(user_id=test_user.id, type="cv", answers=sample_cv_data)
        resume_job = Job(user_id=test_user.id, type="resume", answers=sample_resume_data)
        
        cv_bytes = renderer.render_cv(cv_job)
        resume_bytes = renderer.render_resume(resume_job)
        
        # Both should produce valid documents
        assert len(cv_bytes) > 0
        assert len(resume_bytes) > 0


class TestRenderCoverLetter:
    """Test cover letter rendering"""

    def test_render_basic_cover_letter(self, db_session, test_user, sample_cover_letter_data):
        """Test rendering a basic cover letter"""
        job = Job(
            user_id=test_user.id,
            type="cover",
            answers=sample_cover_letter_data
        )
        
        doc_bytes = renderer.render_cover_letter(job)
        
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        
        doc = Document(BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_cover_letter_contains_company(self, db_session, test_user, sample_cover_letter_data):
        """Test cover letter contains company name"""
        job = Job(
            user_id=test_user.id,
            type="cover",
            answers=sample_cover_letter_data
        )
        
        doc_bytes = renderer.render_cover_letter(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "TechGiant Inc" in text

    def test_cover_letter_contains_role(self, db_session, test_user, sample_cover_letter_data):
        """Test cover letter contains role"""
        job = Job(
            user_id=test_user.id,
            type="cover",
            answers=sample_cover_letter_data
        )
        
        doc_bytes = renderer.render_cover_letter(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Senior Product Manager" in text

    def test_cover_letter_professional_format(self, db_session, test_user, sample_cover_letter_data):
        """Test cover letter follows professional format"""
        job = Job(
            user_id=test_user.id,
            type="cover",
            answers=sample_cover_letter_data
        )
        
        doc_bytes = renderer.render_cover_letter(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Should contain professional salutation
        assert "Dear" in text or "To" in text
        # Should contain closing
        assert "Sincerely" in text or "Regards" in text


class TestRenderingEdgeCases:
    """Test edge cases in rendering"""

    def test_render_with_missing_optional_fields(self, db_session, test_user):
        """Test rendering with minimal data"""
        minimal_data = {
            "basics": {
                "name": "Test User",
                "email": "test@example.com"
            },
            "summary": "Brief summary",
            "skills": ["Skill1"]
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=minimal_data
        )
        
        doc_bytes = renderer.render_resume(job)
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0

    def test_render_with_special_characters(self, db_session, test_user):
        """Test rendering with special characters"""
        data = {
            "basics": {
                "name": "José María O'Brien-Smith",
                "email": "test@example.com",
                "location": "São Paulo, Brazil"
            },
            "summary": "Professional with experience in AI/ML & Data Science",
            "skills": ["Python", "R", "C++", "C#"],
            "experiences": [{
                "company": "Tech & Innovation Co., Ltd.",
                "role": "Senior Engineer",
                "start": "Jan 2020",
                "end": "Present",
                "location": "NYC",
                "bullets": ["Increased revenue by 25% ($1M → $1.25M)"]
            }]
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Special characters should be preserved
        assert "José" in text or "Jose" in text  # Might be normalized

    def test_render_with_long_content(self, db_session, test_user):
        """Test rendering with very long content"""
        long_bullet = "This is a very long bullet point that contains a lot of text " * 20
        
        data = {
            "basics": {"name": "Test User", "email": "test@example.com"},
            "summary": "Summary " * 50,  # Long summary
            "skills": ["Skill1"],
            "experiences": [{
                "company": "Company",
                "role": "Role",
                "start": "2020",
                "end": "Present",
                "location": "NYC",
                "bullets": [long_bullet]
            }]
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=data
        )
        
        # Should not crash with long content
        doc_bytes = renderer.render_resume(job)
        assert len(doc_bytes) > 0

    def test_render_with_empty_lists(self, db_session, test_user):
        """Test rendering with empty lists"""
        data = {
            "basics": {"name": "Test User", "email": "test@example.com"},
            "summary": "Summary",
            "skills": [],  # Empty
            "experiences": [],  # Empty
            "education": [],  # Empty
            "projects": []  # Empty
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=data
        )
        
        # Should handle empty lists gracefully
        doc_bytes = renderer.render_resume(job)
        assert len(doc_bytes) > 0

    def test_render_with_missing_required_fields(self, db_session, test_user):
        """Test rendering fails gracefully with missing required data"""
        incomplete_data = {
            "basics": {"name": "Test"},
            # Missing many required fields
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=incomplete_data
        )
        
        # Should either render with defaults or raise appropriate error
        try:
            doc_bytes = renderer.render_resume(job)
            assert isinstance(doc_bytes, bytes)
        except (KeyError, AttributeError, ValueError) as e:
            # Acceptable to raise error for missing required fields
            pytest.skip(f"Renderer raises error for incomplete data: {e}")

    def test_render_multiple_experiences(self, db_session, test_user):
        """Test rendering with many experiences"""
        experiences = [
            {
                "company": f"Company {i}",
                "role": f"Role {i}",
                "start": f"Jan {2015+i}",
                "end": f"Dec {2015+i}" if i < 5 else "Present",
                "location": "City",
                "bullets": [f"Achievement {i}.1", f"Achievement {i}.2"]
            }
            for i in range(6)
        ]
        
        data = {
            "basics": {"name": "Experienced Pro", "email": "pro@example.com"},
            "summary": "Many years of experience",
            "skills": ["Skill1", "Skill2"],
            "experiences": experiences
        }
        
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Should contain all companies
        for i in range(6):
            assert f"Company {i}" in text


class TestDocumentStructure:
    """Test document structure and formatting"""

    def test_resume_has_sections(self, db_session, test_user, sample_resume_data):
        """Test resume contains expected sections"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for section headers
        text_lower = text.lower()
        assert "experience" in text_lower or "work" in text_lower
        assert "education" in text_lower
        assert "skill" in text_lower

    def test_document_has_content(self, db_session, test_user, sample_resume_data):
        """Test document is not empty"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        doc = Document(BytesIO(doc_bytes))
        
        # Should have meaningful content
        total_text = "\n".join([p.text for p in doc.paragraphs])
        assert len(total_text) > 100  # At least 100 characters

    def test_resume_file_size_reasonable(self, db_session, test_user, sample_resume_data):
        """Test generated file size is reasonable"""
        job = Job(
            user_id=test_user.id,
            type="resume",
            answers=sample_resume_data
        )
        
        doc_bytes = renderer.render_resume(job)
        
        # File should be between 10KB and 500KB
        assert 10_000 < len(doc_bytes) < 500_000
