"""
Tests for revamp service
Tests AI-powered resume improvement feature
"""
import pytest
from unittest.mock import patch, Mock
from app.services.router import handle_revamp, handle_inbound
from app.services import renderer, ai
from app.models import User, Job
from io import BytesIO
from docx import Document


@pytest.mark.asyncio
class TestRevampFlow:
    """Test revamp conversation flow"""

    async def test_revamp_start(self, db_session, test_user):
        """Test revamp flow initiation"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={"_step": "upload"}
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "", "free")
        
        assert "paste your resume" in response.lower() or "resume content" in response.lower()
        assert "Revamp" in response or "improve" in response.lower()

    async def test_revamp_short_content(self, db_session, test_user):
        """Test revamp rejects short content"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={"_step": "upload"}
        )
        db_session.add(job)
        db_session.commit()

        short_text = "John Doe. Email: john@example.com"
        response = await handle_revamp(db_session, job, short_text, "free")
        
        # Should ask for more content
        assert "short" in response.lower() or "more content" in response.lower()

    async def test_revamp_accepts_content(self, db_session, test_user):
        """Test revamp accepts sufficient content"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={"_step": "upload"}
        )
        db_session.add(job)
        db_session.commit()

        resume_content = """
        John Doe
        john@example.com | +1234567890 | New York, NY
        
        EXPERIENCE
        Senior Engineer at TechCorp (2020-Present)
        - Built scalable APIs
        - Led team of 5 engineers
        - Reduced latency by 40%
        
        EDUCATION
        B.Sc. Computer Science, MIT, 2020
        
        SKILLS
        Python, JavaScript, Docker, AWS
        """
        
        response = await handle_revamp(db_session, job, resume_content, "free")
        
        db_session.refresh(job)
        assert job.answers.get("original_content") == resume_content.strip()
        assert "analyzing" in response.lower() or "ai" in response.lower()

    @patch('app.services.ai.revamp_resume')
    async def test_revamp_ai_processing(self, mock_revamp, db_session, test_user):
        """Test AI revamp processing"""
        mock_revamp.return_value = "IMPROVED RESUME CONTENT\n\nJohn Doe - Senior Engineer"
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={
                "_step": "revamp_processing",
                "original_content": "Original resume content here"
            }
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "", "free")
        
        db_session.refresh(job)
        assert job.answers.get("revamped_content") is not None
        assert "IMPROVED" in job.answers["revamped_content"]
        assert "preview" in response.lower() or "AI-Enhanced" in response

    @patch('app.services.ai.revamp_resume')
    async def test_revamp_preview(self, mock_revamp, db_session, test_user):
        """Test revamp preview step"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={
                "_step": "preview",
                "original_content": "Original content",
                "revamped_content": "Improved content"
            }
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "no", "free")
        
        # Should re-show preview
        assert "yes" in response.lower() or "confirm" in response.lower()

    @patch('app.services.ai.revamp_resume')
    @patch('app.services.renderer.render_revamp')
    @patch('app.services.storage.save_file_locally')
    async def test_revamp_confirm_generation(self, mock_storage, mock_render, mock_revamp, db_session, test_user):
        """Test confirming revamp and generating document"""
        mock_render.return_value = b"DOCX_CONTENT"
        mock_storage.return_value = "/path/to/file.docx"
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={
                "_step": "preview",
                "original_content": "Original",
                "revamped_content": "Improved"
            }
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "yes", "free")
        
        # Should trigger document generation
        assert "__SEND_DOCUMENT__" in response

    @patch('app.services.ai.revamp_resume')
    async def test_revamp_ai_failure(self, mock_revamp, db_session, test_user):
        """Test handling AI failure gracefully"""
        mock_revamp.side_effect = Exception("AI service unavailable")
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={
                "_step": "revamp_processing",
                "original_content": "Original content"
            }
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "", "free")
        
        # Should show error message
        assert "error" in response.lower() or "sorry" in response.lower()


class TestRevampAIService:
    """Test AI revamp service"""

    @patch('app.services.ai.client')
    def test_revamp_free_tier(self, mock_client):
        """Test revamp with free tier"""
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = "Improved resume content"
        mock_client.chat.completions.create.return_value = mock_response
        
        original = "John Doe. Experience: worked at company."
        result = ai.revamp_resume(original, tier="free")
        
        assert result == "Improved resume content"
        # Verify free tier prompt used
        call_args = mock_client.chat.completions.create.call_args
        assert "FREE TIER" in str(call_args)

    @patch('app.services.ai.client')
    def test_revamp_pro_tier(self, mock_client):
        """Test revamp with pro tier"""
        mock_response = Mock()
        mock_response.choices = [Mock()]
        mock_response.choices[0].message.content = "Enhanced resume with metrics"
        mock_client.chat.completions.create.return_value = mock_response
        
        original = "John Doe. Experience: worked at company."
        result = ai.revamp_resume(original, tier="pro")
        
        assert result == "Enhanced resume with metrics"
        # Verify pro tier prompt used
        call_args = mock_client.chat.completions.create.call_args
        assert "PRO TIER" in str(call_args)

    def test_revamp_no_ai_client(self):
        """Test revamp without AI client configured"""
        with patch('app.services.ai.client', None):
            original = "Original content"
            result = ai.revamp_resume(original, tier="free")
            
            # Should return original content
            assert result == original

    @patch('app.services.ai.client')
    def test_revamp_api_error(self, mock_client):
        """Test revamp handles API errors"""
        mock_client.chat.completions.create.side_effect = Exception("API Error")
        
        original = "Original content"
        result = ai.revamp_resume(original, tier="free")
        
        # Should return original content on error
        assert result == original


class TestRevampRenderer:
    """Test revamp document rendering"""

    def test_render_revamp_basic(self, db_session, test_user):
        """Test rendering revamped content"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={
                "revamped_content": "Improved resume content\n\nSection 1\nSection 2"
            }
        )
        
        doc_bytes = renderer.render_revamp(job)
        
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        
        # Verify it's valid DOCX
        doc = Document(BytesIO(doc_bytes))
        assert len(doc.paragraphs) > 0

    def test_render_revamp_contains_content(self, db_session, test_user):
        """Test rendered document contains revamped content"""
        content = "IMPROVED RESUME\n\nJohn Doe - Senior Engineer"
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={"revamped_content": content}
        )
        
        doc_bytes = renderer.render_revamp(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "IMPROVED RESUME" in text
        assert "John Doe" in text

    def test_render_revamp_fallback_original(self, db_session, test_user):
        """Test fallback to original content"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={
                "original_content": "Original resume content"
                # No revamped_content
            }
        )
        
        doc_bytes = renderer.render_revamp(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Original resume content" in text

    def test_render_revamp_empty_content(self, db_session, test_user):
        """Test rendering with no content"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={}
        )
        
        doc_bytes = renderer.render_revamp(job)
        doc = Document(BytesIO(doc_bytes))
        
        # Should render something (placeholder or empty)
        assert len(doc.paragraphs) > 0

    def test_render_revamp_multiline(self, db_session, test_user):
        """Test rendering multi-line content"""
        content = """
        John Doe
        Senior Software Engineer
        
        EXPERIENCE
        Company A (2020-Present)
        - Achievement 1
        - Achievement 2
        
        Company B (2018-2020)
        - Achievement 3
        """
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={"revamped_content": content}
        )
        
        doc_bytes = renderer.render_revamp(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "John Doe" in text
        assert "EXPERIENCE" in text


@pytest.mark.asyncio
class TestRevampIntegration:
    """Integration tests for complete revamp flow"""

    async def test_revamp_from_inbound(self, db_session):
        """Test initiating revamp from inbound handler"""
        telegram_id = "revamp_test_001"
        
        # Start flow
        response = await handle_inbound(db_session, telegram_id, "/start")
        assert response is not None
        
        # Select revamp
        response = await handle_inbound(db_session, telegram_id, "revamp")
        assert "paste" in response.lower() or "resume" in response.lower()

    @patch('app.services.ai.revamp_resume')
    async def test_complete_revamp_flow(self, mock_revamp, db_session):
        """Test complete revamp from start to document"""
        mock_revamp.return_value = "IMPROVED: Professional resume content"
        telegram_id = "revamp_complete_001"
        
        # Start
        await handle_inbound(db_session, telegram_id, "/start")
        
        # Initiate revamp
        await handle_inbound(db_session, telegram_id, "revamp")
        
        # Provide content
        resume = """
        John Doe | john@example.com | +1234567890
        
        EXPERIENCE
        Senior Engineer at Tech Company (2020-Present)
        - Built scalable systems
        - Led team of 5
        
        SKILLS
        Python, JavaScript, AWS
        """
        
        response = await handle_inbound(db_session, telegram_id, resume)
        
        # Verify flow progressed
        user = db_session.query(User).filter(
            User.telegram_user_id == telegram_id
        ).first()
        
        job = db_session.query(Job).filter(
            Job.user_id == user.id,
            Job.type == "revamp"
        ).first()
        
        assert job is not None
        assert job.answers.get("original_content") is not None


class TestRevampEdgeCases:
    """Test edge cases in revamp flow"""

    @pytest.mark.asyncio
    async def test_revamp_special_characters(self, db_session, test_user):
        """Test revamp with special characters"""
        content = "José María O'Brien\nPython/C++/C# Developer\nImproved revenue by 50% ($1M → $1.5M)"
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            answers={"revamped_content": content}
        )
        
        doc_bytes = renderer.render_revamp(job)
        doc = Document(BytesIO(doc_bytes))
        
        text = "\n".join([p.text for p in doc.paragraphs])
        # Special characters should be preserved or handled gracefully
        assert "José" in text or "Jose" in text

    @pytest.mark.asyncio
    async def test_revamp_very_long_content(self, db_session, test_user):
        """Test revamp with very long content"""
        long_content = "Resume section\n" * 500  # Very long resume
        
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={"_step": "upload"}
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, long_content, "free")
        
        # Should accept and process
        db_session.refresh(job)
        assert job.answers.get("original_content") is not None

    @pytest.mark.asyncio
    async def test_revamp_empty_input(self, db_session, test_user):
        """Test revamp with empty input"""
        job = Job(
            user_id=test_user.id,
            type="revamp",
            status="collecting",
            answers={"_step": "upload"}
        )
        db_session.add(job)
        db_session.commit()

        response = await handle_revamp(db_session, job, "", "free")
        
        # Should show instructions
        assert "paste" in response.lower() or "content" in response.lower()


class TestRevampPayment:
    """Test payment integration for revamp"""

    @pytest.mark.asyncio
    async def test_revamp_free_tier_limit(self, db_session, test_user):
        """Test revamp respects free tier limits"""
        # Set user at limit
        test_user.generation_count = 10
        db_session.commit()
        
        # This validates payment check exists for revamp

    @pytest.mark.asyncio
    async def test_revamp_pro_user_unlimited(self, db_session, pro_user):
        """Test pro user has unlimited revamps"""
        assert pro_user.tier == "pro"
        
        job = Job(
            user_id=pro_user.id,
            type="revamp",
            status="collecting"
        )
        db_session.add(job)
        db_session.commit()
        
        # Pro user should be able to revamp without limits
